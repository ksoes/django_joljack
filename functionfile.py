from my_settings import engine, my_key

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO
import olefile
import docx
import re
from konlpy.tag import Mecab
import pandas as pd
from tqdm import tqdm
from sklearn.feature_extraction.text import CountVectorizer
import sqlalchemy
import pymysql
import requests
from pandas.io.json import json_normalize
import os
from IPython import display
from newspaper import Article


def connect_sql():  # pymysql / SQLAlchemy
    global conn
    conn = engine.connect()  # mysql 연결, engine은 my_settings.py에서 가져온다
    return engine, conn


# 각 형식별 파일에서 텍스트 추출
def convert_pdf_to_text(url):
    # pdf파일 읽어서 기사 추출
    rsrcmgr = PDFResourceManager()
    retstr=StringIO()
    codec='utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    fp = open(os.getcwd()+url, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password=""
    maxpages=0
    caching=True
    pagenos=set()
    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password, caching=caching, check_extractable=True) :
        interpreter.process_page(page)
    text = retstr.getvalue()
    fp.close()
    device.close()
    retstr.close()
    return text


def convert_hwp_to_text(url):
    f = olefile.OleFileIO(url)   # olefile로 한글 파일 열기
    encoded_text = f.openstream('PrvText').read()  # PrvText 스트림 안의 내용 꺼내기 (유니코드 인코딩 되어 있음)
    decoded_text = encoded_text.decode('UTF-16')   # 유니코드 이므로 UTF-16으로 디코딩
    return decoded_text


def convert_txt_to_text(url) :
    f = open(url)
    text = f.read() # 파일의 내용 전체를 문자열로 리턴
    f.close()
    return text


def convert_docx_to_text(url):
    # .docx형식 파일 읽기
    s=[]
    dox = docx.Document(url)
    for i in dox.paragraphs :
        s.append(i.text)  # 한 줄씩 읽어 와서 리스트 s에 추가
    text = ' '.join(s)  # join메소드로 리스트s의 원소들을 공백을 이용해서 구분
    return text


def convert_link_to_text(url):
  # url읽어와서 본문 가져오기
  article = Article(url, language='ko')
  article.download()  # 기사 다운로드
  article.parse()
  text = article.text
  return text


def clean_sentence(v):
    v = re.sub('인쇄하기|창닫기|인쇄|재배포|기자|무단전재|저작권자', '', v)  # 상단에 쓸모 없는 단어 제거
    # 바이트 변수 제거
    if '\x00' in v :
        v = ' '.join(v.split('\x00'))
    if '\xa0' in v :
        v = ''.join(v.split('\xa0'))
    if '\x0c' in v :
        v = ''.join(v.split('\x0c'))
    # \n, 띄어쓰기 제거
    v = v.replace("\n","")
    v = v.replace(" ","")
    return v


def make_coword():
    """
    동시출현단어쌍 만들기 <일반기사, 가짜기사>
    title_list, stop_word, 기사 본문
    DTM.csv 파일 생성
    DTM을 편리하게 만들어주기 위해 Scikit-Learn에서 제공하는 CountVectorizer를 import 한다.
    자연어 처리(형태소 분석 - 명사) 및 각 단어에 대한 DTM 제작
    타이틀 리스트를 불러와서 title_list 변수에 저장한다.
    """

    t_file_name = open(os.getcwd() + '/etc_file/title_list.txt', 'r', encoding='utf-8')
    title_list = []
    for line in t_file_name.readlines():
        # txt파일을 readlines로 불러오면 개행 문자도 함께 읽어오기 때문에 인덱싱으로 처리해준다.
        title_list.append(line[:])
    t_file_name.close()

    # if s_file_name in locals :
    # 불용어 파일
    s_file_name = open(os.getcwd() + '/etc_file/stop_word.txt', 'r', encoding='utf-8')
    stop_words_list = []
    for line in s_file_name.readlines():
        stop_words_list.append(line.rstrip())
    s_file_name.close()

    # pandas의 read_csv 함수를 이용하여 csv 파일을 불러온다.
    # dataframe누적시킨걸로 csv저장한 파일 -> 추출된 본문 기사.csv
    dataset = pd.read_csv(os.getcwd() + '/etc_file/organized_total_body.csv')  # 닷홈서버

    # 각 형태소별로 분류(Tagging)해주는 Okt 객체를 불러온다.
    mecab = Mecab()

    for title in tqdm(title_list, desc='타이틀 리스트 진행도'):  # title_list에 대해 반복문을 실행
        # 각 타이틀에 대한 6770개 문서의 DTM을 표현하기 위해
        # CountVectorizer 객체를 선언
        cv = CountVectorizer()

        # 각 문서들의 말뭉치(corpus)를 저장할 리스트 선언
        corpus = []

        # 각 타이틀에 대한 문서들의 말 뭉치를 저장한다. (데이터가 많으면 이 부분에서 장시간이 소요될 수 있다.)
        for doc_num in range(len(dataset)):
            # 각 말뭉치에서 명사 리스트를 만든다.
            noun_list = mecab.nouns(dataset[title].loc[doc_num])

            # 이를 문자열로 저장해야하기 때문에 join함수로 공백으로 구분해 corpus에 append한다.
            corpus.append(' '.join(noun_list))

        # CountVectorizer의 fit_transform 함수를 통해 DTM을 한번에 생성할 수 있다.
        DTM_Array = cv.fit_transform(corpus).toarray()

        # feature_names 함수를 사용하면 DTM의 각 열(column)이 어떤 단어에 해당하는지 알 수 있다.
        feature_names = cv.get_feature_names()

        # 추출해낸 데이터를 DataFrame 형식으로 변환한다.
        DTM_DataFrame = pd.DataFrame(DTM_Array, columns=feature_names)

        # if s_file_name in locals :
        # 불용어 열 제거
        DTM_DataFrame.drop(stop_words_list, axis='columns', inplace=True)

        # 최종적으로 DTM을 csv 파일로 저장한다.
        # DTM_DataFrame.to_csv('불용어제거_정리된pdf기사DTM.csv', encoding='utf-8-sig')
        DTM_DataFrame.to_csv('/home/sys1769/mysite/etc_file/_trustnews_DTM.csv', encoding='utf-8-sig')

        # DTM파일 가지고 동시출현단어.csv 파일 만들기
        dataset = pd.read_csv(os.getcwd() + '/etc_file/_trustnews_DTM.csv')

        column_list = dataset.columns[1:]
        word_length = len(column_list)
        count_dict = {}

        for doc_number in tqdm(range(len(dataset)), desc='단어쌍 만들기 진행중'):
            tmp = dataset.loc[doc_number]
        for i, word1 in enumerate(column_list):
            if tmp[word1]:
                for j in range(i + 1, word_length):
                    if tmp[column_list[j]]:
                        count_dict[column_list[i], column_list[j]] = count_dict.get((column_list[i]), 0) + max(
                            tmp[word1], tmp[column_list[j]])

        count_list = []
        for words in count_dict:
            count_list.append([words[0], words[1], count_dict[words]])

        df = pd.DataFrame(count_list, columns=["word1", "word2", "freq"])
        df = df.sort_values(by=['freq'], ascending=False)  # 내림차순
        df = df.reset_index(drop=True)

        # df.to_csv('/home/sys1769/다운로드/불용어제거_정리된pdf기사_networkx.csv', encoding='utf-8-sig')
        df.to_csv('/home/sys1769/mysite/etc_file/_trustnews_networkx.csv', encoding='utf-8-sig')


def make_fcoword():
    # 가짜뉴스_동시출현빈도 파일 생성

    # 타이틀 리스트를 불러와서 title_list 변수에 저장한다.
    t_file_name = open(os.getcwd() + '/etc_file/title_list.txt', 'r', encoding='utf-8')
    title_list = []
    for line in t_file_name.readlines():
        # txt파일을 readlines로 불러오면 개행 문자도 함께 읽어오기 때문에 인덱싱으로 처리해준다.
        title_list.append(line[:])
    t_file_name.close()

    """  
    #if s_file_name in locals :
    # 불용어 파일
    s_file_name = open('/home/sys1769/mysite/etc_file/stop_word.txt', 'r', encoding='utf-8')
    stop_words_list = []
    for line in s_file_name.readlines() :
        stop_words_list.append(line.rstrip())
    s_file_name.close()
    """

    # pandas의 read_csv 함수를 이용하여 csv 파일을 불러온다.
    # dataframe누적시킨걸로 csv저장한 파일 -> 추출된 본문 기사.csv
    dataset = pd.read_csv(os.getcwd() + '/etc_file/fakenewsbody_sum.csv')  # 닷홈서버

    # 각 형태소별로 분류(Tagging)해주는 Okt 객체를 불러온다.
    mecab = Mecab()

    for title in tqdm(title_list, desc='타이틀 리스트 진행도'):  # title_list에 대해 반복문을 실행
        # 각 타이틀에 대한 6770개 문서의 DTM을 표현하기 위해
        # CountVectorizer 객체를 선언
        cv = CountVectorizer()

        # 각 문서들의 말뭉치(corpus)를 저장할 리스트 선언
        corpus = []

        # 각 타이틀에 대한 문서들의 말 뭉치를 저장한다. (데이터가 많으면 이 부분에서 장시간이 소요될 수 있다.)
        for doc_num in range(len(dataset)):
            # 각 말뭉치에서 명사 리스트를 만든다.
            noun_list = mecab.nouns(dataset[title].loc[doc_num])

            # 이를 문자열로 저장해야하기 때문에 join함수로 공백으로 구분해 corpus에 append한다.
            corpus.append(' '.join(noun_list))

        # CountVectorizer의 fit_transform 함수를 통해 DTM을 한번에 생성할 수 있다.
        DTM_Array = cv.fit_transform(corpus).toarray()

        # feature_names 함수를 사용하면 DTM의 각 열(column)이 어떤 단어에 해당하는지 알 수 있다.
        feature_names = cv.get_feature_names()

        # 추출해낸 데이터를 DataFrame 형식으로 변환한다.
        DTM_DataFrame = pd.DataFrame(DTM_Array, columns=feature_names)

        # if s_file_name in locals :
        # 불용어 열 제거
        # DTM_DataFrame.drop(stop_words_list, axis='columns', inplace=True)

        # 최종적으로 DTM을 csv 파일로 저장한다.
        # DTM_DataFrame.to_csv('불용어제거_정리된pdf기사DTM.csv', encoding='utf-8-sig')
        DTM_DataFrame.to_csv('/home/sys1769/mysite/etc_file/fakenewsbody_DTM.csv', encoding='utf-8-sig')

        # DTM파일 가지고 동시출현단어.csv 파일 만들기

        dataset = pd.read_csv(os.getcwd() + '/etc_file/fakenewsbody_DTM.csv')

        column_list = dataset.columns[1:]
        word_length = len(column_list)

        count_dict = {}

        for doc_number in tqdm(range(len(dataset)), desc='단어쌍 만들기 진행중'):
            tmp = dataset.loc[doc_number]
        for i, word1 in enumerate(column_list):
            if tmp[word1]:
                for j in range(i + 1, word_length):
                    if tmp[column_list[j]]:
                        count_dict[column_list[i], column_list[j]] = count_dict.get((column_list[i]), 0) + max(
                            tmp[word1], tmp[column_list[j]])

        count_list = []
        for words in count_dict:
            count_list.append([words[0], words[1], count_dict[words]])

        df = pd.DataFrame(count_list, columns=["word1", "word2", "freq"])
        df = df.sort_values(by=['freq'], ascending=False)  # 내림차순
        df = df.reset_index(drop=True)

        # df.to_csv('/home/sys1769/다운로드/불용어제거_정리된pdf기사_networkx.csv', encoding='utf-8-sig')
        df.to_csv('/home/sys1769/mysite/etc_file/fakenews_networkx.csv', encoding='utf-8-sig')


# db에 빈 테이블 생성
def make_table():
    # 빈 테이블 ex_apidata, ban_apidata, co_occurrence_word 생성
    conn = pymysql.connect(host='localhost', user='user1', password='user1', db='django', charset='utf8')

    sql_ex_apidata = "SET sql_mode=''; create table django.ex_apidata(PRDLST_NM varchar(250) not null, BSSH_NM varchar(150) not null default 'NOPE', primary key(PRDLST_NM,BSSH_NM)) default character set utf8 collate utf8_general_ci;"
    sql_ban_apidata = "SET sql_mode=''; CREATE TABLE django.ban_apidata ( PRDT_NM varchar(250) NOT NULL, MUFC_NM VARCHAR(150) NOT NULL DEFAULT 'NOPE', PRIMARY KEY (PRDT_NM, MUFC_NM) ) default character set utf8 collate utf8_general_ci;"
    sql_coword = "SET sql_mode=''; CREATE TABLE django.co_occurrence_word (word1 varchar(30) not null, word2 varchar(30) not null, count int(11), primary key (word1, word2) ) default character set utf8 collate utf8_general_ci;"
    sql_fakecoword = "SET sql_mode=''; CREATE TABLE django.fake_co_occurrence_word (word1 varchar(30) not null, word2 varchar(30) not null, count int(11), primary key (word1, word2) ) default character set utf8 collate utf8_general_ci;"
    eval_table = "CREATE TABLE django.evaluate_table ( option_gb VARCHAR(10) DEFAULT 'none', eval_text VARCHAR(1500) DEFAULT 'none' ) default character set utf8 collate utf8_general_ci;"

    cur = conn.cursor()
    cur.execute(sql_ex_apidata)
    print(' - api 생성완료')
    cur.execute(sql_ban_apidata)
    print(' - ban 생성완료')
    cur.execute(sql_coword)
    print(' - co 생성완료')
    cur.execute(sql_fakecoword)
    print(' - fco 생성완료')
    cur.execute(eval_table)
    print(' - eval 생성완료')


def check_table():
    # mysql 조회 커리문 - '특정'데이터베이스에 '특정' 테이블이 있는지, 있으면 1, 없으면 0
    engine, conn = connect_sql()
    sql = "SELECT EXISTS ( SELECT 1 FROM Information_schema.tables WHERE table_schema='django' AND table_name='ex_apidata') AS flag"
    df = pd.read_sql(sql, conn)  # 쿼리문 결과를 데이터프레임으로 가져옴
    check = df.iloc[0, 0]  # check 타입은 정수
    return check


# api데이터 받아와서 db에 집어넣기

def put_Rapi():
    # ex_apidata
    # 데이터프레임에 api 전부 받아와 중복제거 후 ex_apidata에 저장

    try:
        total_df = pd.DataFrame(columns=['PRDLST_NM', 'BSSH_NM'])
        i = 1
        j = 1000
        while 1:  # 1. 그냥 1에서 수정
            print("api받아오는중...")
            url = "http://openapi.foodsafetykorea.go.kr/api/"+my_key+"/I0030/json/" + str(i) + "/" + str(j)
            data = requests.get(url).json()
            # print("data['I0030']['RESULT']['MSG'] = ",data['I0030']['RESULT']['MSG']) # right="정상처리되었습니다."
            if data['I0030']['RESULT']['MSG'] != "정상처리되었습니다.":
                print('데이터가 더 이상 없으므로 받아오기 종료')
                break
            else:
                body = [data['I0030']['row']]
                a = pd.json_normalize(data['I0030']['row'])
                # BSSH_NM : 업소_명, # PRDLST_NM : 품목_명
                info_df = a[['PRDLST_NM', 'BSSH_NM']]
                # print(info_df) # dataFrame
                info_df.columns = ['PRDLST_NM', 'BSSH_NM']
                total_df = pd.concat([total_df, info_df], axis=0, ignore_index=True)
                i += 1000
                j += 1000
                print(total_df)

        print('while문 탈출!')
        # 중복 데이터 제거
        print('중복데이터 제거합니다')
        total_df = total_df.drop_duplicates(['PRDLST_NM', 'BSSH_NM'], keep='first')  # first:첫번째행남김, last:마지막행남김
        # 제품명 열 데이터 띄어쓰기 제거
        print('제품명 열 띄어쓰기 제거합니다')
        total_df['PRDLST_NM'] = total_df['PRDLST_NM'].str.replace(' ', '')
        total_df['BSSH_NM'] = total_df['BSSH_NM'].str.replace(' ', '')

    except:
        total_df = pd.read_csv(os.getcwd() + '/etc_file/ex_apidata.csv')
        print("api - 다운로드 파일로 생성중")
        del total_df['Unnamed: 2']

    # sql에 저장하기
    print('sql에 저장합니다')
    total_df.to_sql(name='ex_apidata', con=engine, if_exists='replace', index=False)
    print('저장 끝')
    return total_df


def put_Bapi():
    # ban_apidata
    # 데이터프레임에 api 전부 받아와 중복제거 후 ban_apidata에 저장
    try:
        bantotal_df = pd.DataFrame(columns=['PRDT_NM', 'MUFC_NM'])
        i = 1
        j = 1000
        while 1:
            url = "http://openapi.foodsafetykorea.go.kr/api/"+my_key+"/I2715/json/" + str(i) + "/" + str(j)
            data = requests.get(url).json()
            if data['I2715']['RESULT']['MSG'] == "해당하는 데이터가 없습니다.":
                break
            else:
                body = [data['I2715']['row']]
                a = pd.json_normalize(data['I2715']['row'])
                info_df = a[['PRDT_NM', 'MUFC_NM']]
                info_df.columns = ['PRDT_NM', 'MUFC_NM']
                bantotal_df = pd.concat([bantotal_df, info_df], axis=0, ignore_index=True)
                i += 1000
                j += 1000

        # 위해, 판매중지 식품
        fake_df = pd.DataFrame(columns=['PRDT_NM', 'MUFC_NM'])
        i = 1
        j = 1000
        while tqdm(1):
            url = "http://openapi.foodsafetykorea.go.kr/api/"+my_key+"/I0490/json/" + str(i) + "/" + str(j)
            data = requests.get(url).json()
            if data['I0490']['RESULT']['MSG'] == "해당하는 데이터가 없습니다.":
                break
            else:
                body = [data['I0490']['row']]
                b = pd.json_normalize(data['I0490']['row'])
                fake_info_df = b[['PRDTNM', 'BSSHNM']]
                fake_info_df.columns = ['PRDT_NM', 'MUFC_NM']
                fake_df = pd.concat([fake_df, fake_info_df], axis=0, ignore_index=True)
                i += 1000
                j += 1000

        # 해외위해식품df 위해.판매중지df 합치기
        bantotal_df = pd.concat([bantotal_df, fake_df], axis=0, ignore_index=True)
        # 중복값 제거
        bantotal_df = bantotal_df.drop_duplicates(['PRDT_NM', 'MUFC_NM'], keep=False)
        bantotal_df = bantotal_df.reset_index(drop=True)
        # 제품명 열 데이터 띄어쓰기 제거
        bantotal_df['PRDT_NM'] = bantotal_df['PRDT_NM'].str.replace(' ', '')
    except:
        print('ban - 다운로드 파일로 생성중..')
        bantotal_df = pd.read_csv(os.getcwd() + '/etc_file/ban_apidata.csv')
        del bantotal_df['Unnamed: 2']

    bantotal_df.to_sql(name='ban_apidata', con=engine, if_exists='replace', index=False)  # index=First&last 하면 오류남..
    return bantotal_df


def put_Cword():
    # co_occurrence_word
    make_coword()  # 동출빈.csv파일을 가져와야되니 우선 파일을 생성
    coword_df = pd.read_csv(os.getcwd() + '/etc_file/_trustnews_networkx.csv')
    del coword_df['Unnamed: 0']  # 첫번째 쓸데없는 열 삭제
    coword_df.rename(columns={'freq': 'count'}, inplace=True)  # 열이름 변경
    coword_df.to_sql(name='co_occurrence_word', con=engine, if_exists='append', index=False)
    return coword_df


def put_FCword():  # fake_co_occurrence_word
    make_fcoword()  # fake동출빈.csv파일을 가져와야되니 우선 파일을 생성
    fcoword_df = pd.read_csv(os.getcwd() + '/etc_file/fakenews_networkx.csv')
    del fcoword_df['Unnamed: 0']
    fcoword_df.rename(columns={'freq': 'count'}, inplace=True)  # 열이름 변경
    fcoword_df.to_sql(name='fake_co_occurrence_word', con=engine, if_exists='append', index=False)
    return fcoword_df


def check_presence_data():  # 테이블에 데이터가 있는지 확인

    sql_api = "select count(*) from ex_apidata;"
    cf = pd.read_sql(sql_api, conn)  # 쿼리문 결과를 데이터프레임으로 가져옴
    sqlapi = cf.iloc[0, 0]  # check 타입은 정수

    sql_ban = "select count(*) from ban_apidata;"
    cf = pd.read_sql(sql_ban, conn)
    sqlban = cf.iloc[0, 0]

    sql_co = "select count(*) from co_occurrence_word"
    cf = pd.read_sql(sql_co, conn)
    sqlco = cf.iloc[0, 0]

    sql_fco = "select count(*) from fake_co_occurrence_word"
    cf = pd.read_sql(sql_fco, conn)
    sqlfco = cf.iloc[0, 0]

    return sqlapi, sqlban, sqlco, sqlfco


def take_Rapi():  # mysql에서 ex_apidata 테이블(데이터) 가져오기
    total_df = pd.read_sql_table('ex_apidata', conn)
    # total_df = pd.read_csv(os.getcwd() + '/etc_file/ex_apidata.csv')
    return total_df


def take_Bapi():  # mysql에서 ban_apidata 테이블(데이터) 가져오기
    bantotal_df1 = pd.read_sql_table('ban_apidata', conn)
    # bantotal_df1 = pd.read_csv(os.getcwd() + '/etc_file/ban_apidata.csv')
    return bantotal_df1


def take_Cword():  # mysql에서 co_occurrence_word 테이블(데이터) 가져오기
    coword_df = pd.read_sql_table('co_occurrence_word', conn)
    # coword_df = pd.read_csv(os.getcwd() + '/etc_file/_trustnews_networkx.csv')
    # del coword_df['Unnamed: 0']  # csv파일에서 불러오면 해당 항목이 있어서 삭제해줘야함, mysql은 해당사항 없음
    return coword_df


def take_FCword():  # 계산된 fake동출빈파일.csv 데이터 가져오기
    fcoword_df = pd.read_sql_table('fake_co_occurrence_word', conn)
    # fcoword_df = pd.read_csv(os.getcwd() + '/etc_file/fakenews_networkx.csv')
    # del fcoword_df['Unnamed: 0']  # csv파일에서 불러오면 해당 항목이 있어서 삭제해줘야함, mysql은 해당사항 없음
    # conn.close() # sql 마지막 연결
    return fcoword_df


# 허위과대광고문구 있는지 조회하는 함수
def fake_text_find(v):
    # 허위과대문구 텍스트파일에서 리스트로 불러오기
    fake_ad_file = open(os.getcwd() + "/etc_file/fake_sentence.txt", 'r')  # 로컬서버
    fake_ad_list = []
    while True:
        fake_ad = fake_ad_file.readline()
        if not fake_ad:
            break
        fake_ad = fake_ad.strip()
        fake_ad_list.append(fake_ad)
    fake_ad_file.close()
    # fake_ad_list # type = list
    # 허위과대문구 본문에 있는지 조회
    fake_ad_count = 0
    # print('')
    for i in range(len(fake_ad_list)):
        if fake_ad_list[i] in v:
            print("삑 : {0}".format(fake_ad_list[i]))
            fake_ad_count += 1
    return fake_ad_count


def get_recommendations(article, cosine_sim, indices, data_df):
    idx = indices[article]  # 입력한 기사로 인덱스를 받아옴
    sim_scores = list(enumerate(cosine_sim[idx]))  # 모든 기사에 대해 입력한 기사와의 유사도를 구함
    sim_scores = sorted(sim_scores, key=lambda x: x[1], reverse=True)  # 유사도에 따라 기사들을 정렬
    sim_scores = sim_scores[1:6]  # 가장 유사한 기사 5개를 받아옴
    movie_indices = [i[0] for i in sim_scores]  # 가장 유사한 기사 5개의 인덱스를 받아옴
    return data_df['code'].iloc[movie_indices]  # 가장 유사한 기사의 신뢰도(높or낮)를 리턴

